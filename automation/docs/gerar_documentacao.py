#!/usr/bin/env python3
"""
Gera o documento Word com toda a documentacao da automacao.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def criar_documento() -> Path:
    doc = Document()

    # ===================================================================
    # ESTILOS
    # ===================================================================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titulo principal
    titulo = doc.add_heading('DOCUMENTACAO DA AUTOMACAO', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run('Programa Mais Top - Download de Bases\n')
    run.font.size = Pt(14)
    run.font.bold = True
    run2 = subtitulo.add_run('Cadastro e Acessos')
    run2.font.size = Pt(12)
    run2.font.italic = True

    doc.add_paragraph()
    doc.add_paragraph('Data: 03/06/2026')
    doc.add_paragraph('Versao: 1.0')
    doc.add_page_break()

    # ===================================================================
    # 1. O QUE E ESTA AUTOMACAO
    # ===================================================================
    doc.add_heading('1. O QUE E ESTA AUTOMACAO?', level=1)

    doc.add_paragraph(
        'Esta automacao foi criada para eliminar o trabalho manual e repetitivo '
        'de baixar as planilhas de cadastro e acessos do site do Programa Mais Top.'
    )

    doc.add_heading('1.1. Problema Atual (Manual)', level=2)
    p = doc.add_paragraph()
    p.add_run('Hoje voce faz isso todo dia:').bold = True
    passos = [
        '1. Abre o navegador',
        '2. Acessa programamaistop.com.br',
        '3. Digita usuario e senha',
        '4. Clica em "Relatorios"',
        '5. Clica em "Download Cadastro" e espera',
        '6. Clica em "Download Acessos" e espera',
        '7. Move os arquivos para a pasta correta (bases_recorrentes/DDMMYY/)',
        '8. Abre o Power BI Desktop',
        '9. Atualiza as bases no Power Query',
        '10. Salva e publica',
        '11. Tira print e manda no Teams',
    ]
    for passo in passos:
        doc.add_paragraph(passo, style='List Number')

    doc.add_heading('1.2. Solucao Proposta (Automatizada)', level=2)
    p = doc.add_paragraph()
    p.add_run('Com o script, o computador faz sozinho:').bold = True
    auto_passos = [
        '1. Abre o navegador automaticamente',
        '2. Preenche login e senha',
        '3. Navega ate Relatorios',
        '4. Baixa Cadastro e Acessos',
        '5. Salva na pasta correta com data e hora',
    ]
    for passo in auto_passos:
        doc.add_paragraph(passo, style='List Number')

    p = doc.add_paragraph()
    run = p.add_run('Voce so precisa fazer: ')
    run.bold = True
    p.add_run('atualizar no Power BI Desktop, publicar e enviar no Teams.')

    doc.add_page_break()

    # ===================================================================
    # 2. COMO FUNCIONA O SCRIPT (PASSO A PASSO VISUAL)
    # ===================================================================
    doc.add_heading('2. COMO FUNCIONA O SCRIPT?', level=1)

    doc.add_paragraph(
        'O script e escrito em Python e usa uma biblioteca chamada Playwright. '
        'Pense no Playwright como um "robo que controla o navegador".'
    )

    doc.add_heading('2.1. Fluxo Visual do Script', level=2)

    fluxo = [
        ('INICIO', 'O script comeca a rodar'),
        ('PASSO 1: Abrir Navegador', 'Playwright abre um Chrome/Chromium. Pode ser invisivel (headless) ou visivel.'),
        ('PASSO 2: Acessar o Site', 'Navegador vai para programamaistop.com.br'),
        ('PASSO 3: Login', 'Script preenche o campo usuario, depois senha, depois clica em Entrar'),
        ('PASSO 4: Navegar', 'Script procura e clica no menu "Relatorios"'),
        ('PASSO 5: Download Cadastro', 'Script procura o botao/link de Cadastro e clica. O arquivo e baixado.'),
        ('PASSO 6: Download Acessos', 'Mesmo processo para Acessos.'),
        ('PASSO 7: Organizar Arquivos', 'Script move os arquivos baixados para bases_recorrentes/DDMMYY/'),
        ('PASSO 8: Fechar Navegador', 'Script encerra e mostra resumo.'),
        ('FIM', 'Voce recebe os dois arquivos na pasta certa.'),
    ]

    for titulo_passo, desc in fluxo:
        p = doc.add_paragraph()
        run = p.add_run(titulo_passo + ': ')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)
        p.add_run(desc)

    doc.add_heading('2.2. Diagrama Simplificado', level=2)
    diagrama = """
    +------------------+
    |  Rodar Script    |
    +--------+---------+
             |
             v
    +--------+---------+
    | Abrir Navegador  |<------ Chrome/Chromium controlado pelo script
    +--------+---------+
             |
             v
    +--------+---------+
    | Acessar Site     |<------ programamaistop.com.br
    +--------+---------+
             |
             v
    +--------+---------+
    |  Fazer Login     |<------ Preenche usuario + senha automaticamente
    +--------+---------+
             |
             v
    +--------+---------+
    |  Relatorios      |<------ Clica no menu de relatorios
    +--------+---------+
             |
       +-----+-----+
       |           |
       v           v
+------+-----+ +---+---------+
| Download   | | Download    |
| Cadastro   | | Acessos     |
+------+-----+ +---+---------+
       |           |
       +-----+-----+
             |
             v
    +--------+---------+
    | Salvar em        |<------ bases_recorrentes/030626/
    | bases_recorrentes|
    +--------+---------+
             |
             v
    +--------+---------+
    |    FIM           |
    +------------------+
    """
    p = doc.add_paragraph()
    p.add_run(diagrama).font.name = 'Consolas'

    doc.add_page_break()

    # ===================================================================
    # 3. O QUE SAO SELETORES CSS? (EXPLICACAO DIDATICA)
    # ===================================================================
    doc.add_heading('3. O QUE SAO SELETORES CSS?', level=1)

    doc.add_paragraph(
        'Esta e a parte MAIS IMPORTANTE para entender. Sem os seletores corretos, '
        'o script nao funciona. Mas e simples de ajustar!'
    )

    doc.add_heading('3.1. Analogia Simples', level=2)
    doc.add_paragraph(
        'Imagine que o site e uma casa e cada botao/campo e uma porta ou gaveta. '
        'O seletor CSS e como um "endereco" que diz ao robo EXATAMENTE onde clicar.'
    )

    p = doc.add_paragraph()
    p.add_run('Exemplo do mundo real:\n').bold = True
    p.add_run(
        'Se eu te disser "vai na cozinha, abre a gaveta de baixo da pia, '
        'pega a colher de pau", voce sabe exatamente onde ir. '
        'O seletor CSS faz a mesma coisa, mas para elementos na pagina web.'
    )

    doc.add_heading('3.2. Como um Seletor CSS Funciona', level=2)

    doc.add_paragraph('Todo site e feito de HTML. Um campo de login no HTML pode ser assim:')

    exemplo_html = '''
    <input type="text" name="username" id="user" placeholder="Digite seu usuario">
    <input type="password" name="password" id="pass" placeholder="Digite sua senha">
    <button type="submit" class="btn-login">Entrar</button>
    '''
    p = doc.add_paragraph()
    p.add_run(exemplo_html).font.name = 'Consolas'
    p.add_run('\n(Fonte: Consolas = codigo)')

    doc.add_paragraph('O script precisa encontrar esses elementos. Os seletores possiveis sao:')

    tabela = doc.add_table(rows=1, cols=3)
    tabela.style = 'Light Grid Accent 1'
    hdr_cells = tabela.rows[0].cells
    hdr_cells[0].text = 'Elemento'
    hdr_cells[1].text = 'Seletor CSS'
    hdr_cells[2].text = 'Significado'

    linhas = [
        ('Campo usuario', 'input[name="username"]', 'Encontra <input> com atributo name="username"'),
        ('Campo senha', 'input[type="password"]', 'Encontra <input> com type="password"'),
        ('Botao entrar', 'button[type="submit"]', 'Encontra <button> com type="submit"'),
        ('Botao entrar (2)', 'button:has-text("Entrar")', 'Encontra <button> que contem o texto "Entrar"'),
        ('Por ID', '#user', 'Encontra elemento com id="user"'),
        ('Por classe', '.btn-login', 'Encontra elemento com class="btn-login"'),
    ]

    for elem, seletor, sig in linhas:
        row = tabela.add_row().cells
        row[0].text = elem
        row[1].text = seletor
        row[2].text = sig

    doc.add_heading('3.3. Como Descobrir os Seletores do Site Real', level=2)

    instrucoes = [
        '1. Abra o Chrome e acesse programamaistop.com.br',
        '2. Aperte F12 (ou Ctrl+Shift+I) para abrir as Ferramentas do Desenvolvedor',
        '3. Clique no icone de selecao (canto superior esquerdo, parece um cursor num quadrado) ou aperte Ctrl+Shift+C',
        '4. Passe o mouse sobre o campo de usuario na pagina',
        '5. Clique no campo de usuario',
        '6. A janela do F12 vai mostrar o codigo HTML do elemento destacado em azul',
        '7. Clique com o botao DIREITO no elemento azul no painel do F12',
        '8. Clique em "Copy" > "Copy selector" (ou "Copiar" > "Copiar seletor")',
        '9. Cole esse seletor no script download_bases.py',
        '10. Repita para: campo senha, botao login, menu relatorios, botao cadastro, botao acessos',
    ]

    for instr in instrucoes:
        doc.add_paragraph(instr, style='List Number')

    doc.add_heading('3.4. Exemplo Visual', level=2)

    exemplo_visual = """
    +------------------------------------------+
    |  programamaistop.com.br/login            |
    +------------------------------------------+
    |                                          |
    |   [____________________]  <-- campo      |
    |    Digite seu usuario       usuario      |
    |                                          |
    |   [____________________]  <-- campo      |
    |    Digite sua senha         senha        |
    |                                          |
    |   [      ENTRAR      ]    <-- botao     |
    |                                          |
    +------------------------------------------+

    Seletor do campo usuario (exemplo):
    input[placeholder*="usuário"]

    Seletor do campo senha (exemplo):
    input[type="password"]

    Seletor do botao (exemplo):
    button:has-text("Entrar")
    """
    p = doc.add_paragraph()
    p.add_run(exemplo_visual).font.name = 'Consolas'

    doc.add_page_break()

    # ===================================================================
    # 4. APLICACOES GRAFICAS UTILIZAVEIS
    # ===================================================================
    doc.add_heading('4. APLICACOES GRAFICAS UTILIZAVEIS', level=1)

    doc.add_paragraph(
        'Alem do script Python, existem ferramentas visuais (com interface grafica) '
        'que tambem podem automatizar esse tipo de tarefa. Aqui estao as principais:'
    )

    ferramentas = [
        ('Power Automate Desktop', 'Microsoft', 'Gratuito (incluso no Windows 10/11). Interface visual de arrastar e soltar. Grava acoes do mouse/teclado. Nao precisa saber programar.'),
        ('UiPath Community', 'UiPath', 'Gratuito para uso pessoal. Muito poderoso. Interface visual. Curva de aprendizado media.'),
        ('Automation Anywhere', 'Automation Anywhere', 'Pago. Enterprise. Para empresas grandes.'),
        ('Playwright (modo visivel)', 'Microsoft', 'O que usamos no script. Pode rodar COM navegador aberto (modo --visible) para voce acompanhar.'),
        ('Selenium IDE', 'Open Source', 'Extensao do Chrome/Firefox. Grava cliques e replay. Muito simples.'),
        ('Macro Recorder (Windows)', 'Varios', 'Grava movimentos de mouse e teclado e repete. Solucao mais simples possivel.'),
    ]

    for nome, empresa, desc in ferramentas:
        doc.add_heading(f'4.{ferramentas.index((nome, empresa, desc)) + 1}. {nome}', level=2)
        p = doc.add_paragraph()
        p.add_run(f'Empresa: ').bold = True
        p.add_run(f'{empresa}\n')
        p.add_run(f'Descricao: ').bold = True
        p.add_run(desc)

    doc.add_heading('4.7. Qual a Melhor Opcao Para Voce?', level=2)

    p = doc.add_paragraph()
    p.add_run('Recomendacao: ').bold = True
    p.add_run(
        'Se voce nao sabe programar, o '
    )
    run = p.add_run('Power Automate Desktop')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(
        ' e a melhor escolha. Ele ja vem instalado no Windows 11 '
        '(e pode ser baixado gratis no Windows 10). Voce apenas '
        'clica em "Gravar" e faz as acoes uma vez. Depois ele repete sozinho.'
    )

    doc.add_page_break()

    # ===================================================================
    # 5. POWERSHELL + POWER BI
    # ===================================================================
    doc.add_heading('5. COMO AUTOMATIZAR O POWER BI COM POWERSHELL?', level=1)

    doc.add_paragraph(
        'O Power BI Desktop e um programa de computador (aplicacao desktop), '
        'entao nao da para automatizar pelo navegador. Mas existe uma forma '
        'usando PowerShell (linguagem de script do Windows).'
    )

    doc.add_heading('5.1. O que e PowerShell?', level=2)
    doc.add_paragraph(
        'PowerShell e como um "Prompt de Comando avancado" do Windows. '
        'Com ele voce pode dar comandos para o computador fazer coisas, '
        'inclusive abrir programas e clicar em botoes (usando uma biblioteca '
        'chamada AutoIt ou UIAutomation).'
    )

    doc.add_heading('5.2. Limitacoes Importantes', level=2)

    limitacoes = [
        'O Power BI Desktop NAO tem API oficial para automatizar cliques internos.',
        'Atualizar fontes de dados no Power Query requer interface grafica.',
        'Publicar online (Publish) requer login na conta Microsoft no Power BI Desktop.',
        'Automatizar isso com PowerShell e possivel, mas FRAGIL (quebra facil se a tela mudar).',
    ]
    for lim in limitacoes:
        doc.add_paragraph(lim, style='List Bullet')

    doc.add_heading('5.3. O Que PODERIA Ser Feito Com PowerShell', level=2)

    p = doc.add_paragraph()
    p.add_run('Ideia de script PowerShell (conceitual):').italic = True

    ps_exemplo = (
        '# 1. Abre o Power BI Desktop com o arquivo\n'
        'Start-Process "C:\\\\Program Files\\\\Microsoft Power BI Desktop\\\\bin\\\\PBIDesktop.exe" `\n'
        '    -ArgumentList "C:\\\\Projetos\\\\Programa_mais_top\\\\power_bi\\\\Dashboard_TOP_V7.pbix"\n\n'
        '# 2. Aguarda o Power BI abrir\n'
        'Start-Sleep -Seconds 15\n\n'
        '# 3. Usa AutoIt para clicar em "Transformar Dados"\n'
        '# (requer biblioteca adicional)\n'
        '# ... codigo de automacao de interface ...\n\n'
        '# 4. Atualiza as fontes de dados\n'
        '# ... codigo ...\n\n'
        '# 5. Salva e fecha\n'
        '# ... codigo ...\n'
    )
    p = doc.add_paragraph()
    p.add_run(ps_exemplo).font.name = 'Consolas'

    doc.add_heading('5.4. Conclusao sobre Power BI', level=2)
    p = doc.add_paragraph()
    run = p.add_run('NAO E RECOMENDADO automatizar o Power BI Desktop. ')
    run.bold = True
    run.font.color.rgb = RGBColor(192, 0, 0)
    p.add_run(
        'A Microsoft nao oferece suporte oficial para isso e e muito instavel. '
        'O ideal e manter a parte do Power BI manual (apenas 3 cliques no Transformar Dados).'
    )

    doc.add_page_break()

    # ===================================================================
    # 6. GUIA DE INSTALACAO
    # ===================================================================
    doc.add_heading('6. GUIA DE INSTALACAO DO SCRIPT', level=1)

    doc.add_heading('6.1. Requisitos', level=2)
    reqs = [
        'Python 3.9 ou superior instalado',
        'Conexao com internet',
        'Credenciais do site programamaistop.com.br',
        'Permissao do gestor (seguranca da informacao)',
    ]
    for r in reqs:
        doc.add_paragraph(r, style='List Bullet')

    doc.add_heading('6.2. Passos de Instalacao', level=2)

    instalacao = [
        ('Abra o Terminal', 'No VS Code, aperte Ctrl+` (crase) ou abra o CMD/PowerShell'),
        ('Va para a pasta', 'cd C:\\\\caminho\\\\para\\\\Programa_mais_top\\\\automation'),
        ('Instale dependencias', 'python setup.py'),
        ('Configure credenciais', 'Copie .env.example para .env e edite'),
        ('Ajuste os seletores', 'Abra o site com F12 e copie os seletores reais'),
        ('Teste visivel', 'python download_bases.py --visible --screenshot'),
        ('Teste automatico', 'python download_bases.py'),
        ('Agende no cron/Windows', 'Configure para rodar automaticamente'),
    ]

    for i, (passo, cmd) in enumerate(instalacao, 1):
        p = doc.add_paragraph()
        p.add_run(f'Passo {i}: {passo}\n').bold = True
        p.add_run(cmd).font.name = 'Consolas'

    doc.add_heading('6.3. Comandos Uteis', level=2)

    tabela_cmds = doc.add_table(rows=1, cols=3)
    tabela_cmds.style = 'Light Grid Accent 1'
    hdr = tabela_cmds.rows[0].cells
    hdr[0].text = 'Comando'
    hdr[1].text = 'O que faz'
    hdr[2].text = 'Quando usar'

    cmds = [
        ('python download_bases.py', 'Roda em modo invisivel', 'Producao - dia a dia'),
        ('python download_bases.py --visible', 'Roda com navegador aberto', 'Teste / Debug'),
        ('python download_bases.py --screenshot', 'Salva fotos de cada etapa', 'Teste / Debug'),
        ('crontab -e', 'Edita agendamento Linux', 'Automacao diaria'),
    ]

    for cmd, faz, quando in cmds:
        row = tabela_cmds.add_row().cells
        row[0].text = cmd
        row[1].text = faz
        row[2].text = quando

    doc.add_page_break()

    # ===================================================================
    # 7. CHECKLIST DE APROVACAO DO GESTOR
    # ===================================================================
    doc.add_heading('7. CHECKLIST PARA DISCUTIR COM O GESTOR', level=1)

    doc.add_paragraph(
        'Antes de colocar a automacao em producao, discuta estes pontos '
        'com seu gestor ou equipe de TI/Seguranca:'
    )

    checklist = [
        ('Seguranca de Credenciais', 'As credenciais ficam em um arquivo .env. Nao e o mais seguro. Ideal seria usar um cofre de senhas (Azure Key Vault, HashiCorp Vault) ou autenticacao SSO.'),
        ('Politica de Acesso', 'O script acessa o site corporativo automaticamente. Isso viola alguma politica de uso aceitavel (AUP) da empresa?'),
        ('Logs e Rastreabilidade', 'O site pode registrar que um "bot" acessou. Isso pode ser confundido com ataque. O script usa User-Agent real do Chrome para parecer humano.'),
        ('Protecao Anti-Bot', 'O site pode ter Cloudflare, reCAPTCHA ou WAF que bloqueia acessos automatizados. Testar primeiro em ambiente controlado.'),
        ('Backup Manual', 'Manter o processo manual como fallback caso a automacao falhe.'),
        ('Acesso a Dados', 'O script baixa dados de cadastro (possivelmente PII/dados pessoais). Isso esta em conformidade com a LGPD/politica de dados da empresa?'),
        ('Aprovacao Formal', 'Obter aprovacao por escrito (e-mail) do gestor antes de implementar.'),
    ]

    for i, (titulo, desc) in enumerate(checklist, 1):
        doc.add_heading(f'7.{i}. {titulo}', level=2)
        doc.add_paragraph(desc)

    doc.add_page_break()

    # ===================================================================
    # 8. PROXIMOS PASSOS
    # ===================================================================
    doc.add_heading('8. PROXIMOS PASSOS RECOMENDADOS', level=1)

    proximos = [
        '1. Leia este documento completo',
        '2. Mostre para seu gestor e discuta o checklist do item 7',
        '3. Se aprovado, faca um teste em ambiente controlado',
        '4. Descubra os seletores CSS reais do site (item 3.3)',
        '5. Configure as credenciais no arquivo .env',
        '6. Execute o script em modo visivel para validar',
        '7. Ajuste os seletores se necessario',
        '8. Execute em modo headless (automatico)',
        '9. Agende para rodar diariamente',
        '10. Monitore os logs nas primeiras semanas',
    ]

    for passo in proximos:
        doc.add_paragraph(passo, style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('--- FIM DO DOCUMENTO ---')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Salva
    output_path = Path(__file__).resolve().parent / 'Documentacao_Automacao_MaisTop.docx'
    doc.save(str(output_path))
    return output_path


if __name__ == '__main__':
    caminho = criar_documento()
    print(f'Documento gerado: {caminho}')
